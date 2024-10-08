from flask import Flask, render_template, request, jsonify, redirect, url_for, session, flash
from functools import wraps
from flask_mysqldb import MySQL
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField
from wtforms.validators import InputRequired, Length, EqualTo, Email, Regexp
from email_validator import validate_email, EmailNotValidError 
from werkzeug.security import generate_password_hash, check_password_hash
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics import jaccard_score
from transformers import BertTokenizer, BertModel
import torch
import fitz  # PyMuPDF
import docx
import os
import chardet
from langdetect import detect, DetectorFactory
import difflib
 

# Create the Flask application instance and specify the main templates folder
app = Flask(__name__)

# Common configuration
app.config['SECRET_KEY'] = 'sjhjshfjwgsdhgsdyegwdfyhuerbfhjjew'
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''
app.config['MYSQL_DB'] = 'documentsimilarity'
app.config['MYSQL_CURSORCLASS'] = 'DictCursor'

# Initialize MySQL
mysql = MySQL(app)

# Initialize BERT tokenizer and models
tokenizer_single = BertTokenizer.from_pretrained('bert-base-uncased')
model_single_path = 'Custom BERT model/bert_entire_model.pth'  # Path to your custom model
model_single = torch.load(model_single_path, map_location=torch.device('cpu'))
model_single.eval()
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model_single.to(device)

# Load the local BERT model and tokenizer
tokenizer_double = BertTokenizer.from_pretrained('./bert_model')
model_double = BertModel.from_pretrained('./bert_model')
device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
model_double.to(device)

def get_binary_representation(text, tokenizer):
    tokens = tokenizer.tokenize(text)
    unique_tokens = set(tokens)
    return unique_tokens

# Function to generate BERT embeddings for single document comparison
def get_bert_embedding_single(text, tokenizer, model, device):
    inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True, max_length=512).to(device)
    with torch.no_grad():
        outputs = model(**inputs)
    return outputs.last_hidden_state.mean(dim=1).squeeze()

# Function to generate BERT embeddings for double document comparison
def get_bert_embedding_double(text, tokenizer, model, device):
    inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True).to(device)
    with torch.no_grad():
        outputs = model(**inputs)
    return outputs.last_hidden_state.mean(dim=1).squeeze()

# Function to detect file encoding
def detect_encoding(file_path):
    with open(file_path, 'rb') as f:
        raw_data = f.read(10000)
        result = chardet.detect(raw_data)
    return result['encoding']

# Function to load existing documents
def load_existing_documents(dataset_paths):
    documents = []
    for path in dataset_paths:
        try:
            encoding = detect_encoding(path)
            print(f"Loading document: {path} with encoding: {encoding}")  # Debugging line
            with open(path, 'r', encoding=encoding) as f:
                documents.append((os.path.basename(path), f.read()))
                print(f"Successfully loaded document: {path}")  # Debugging line
        except Exception as e:
            print(f"Error reading file {path} with detected encoding: {encoding}. Trying other encodings...")
            # Try with common encodings
            encodings = ['utf-8', 'latin-1', 'windows-1252']
            for enc in encodings:
                try:
                    with open(path, 'r', encoding=enc) as f:
                        documents.append((os.path.basename(path), f.read()))
                        print(f"Successfully loaded document: {path} with encoding: {enc}")  # Debugging line
                        break
                except Exception as inner_e:
                    print(f"Error reading file {path} with encoding {enc}: {str(inner_e)}")
            else:
                print(f"Failed to load document: {path} with all attempted encodings")
    return documents

# Function to check allowed file types
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx', 'txt'}

def read_text_from_file(file, filename):
    """Read and extract text from the uploaded file."""
    ext = filename.rsplit('.', 1)[1].lower()
    text = None
    error = None

    try:
        if ext == 'txt':
            encodings = ['utf-8', 'latin1', 'windows-1252']
            for encoding in encodings:
                try:
                    text = file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if text is None:
                error = "Unable to decode text file."
        
        elif ext == 'pdf':
            import fitz  # PyMuPDF
            file.seek(0)  # Reset file pointer to the start
            doc = fitz.open(stream=file.read(), filetype='pdf')
            text = ""
            for page in doc:
                text += page.get_text()
            # Check for images in the document
            if any(img for img in doc.get_page_images(0, full=True)):
                error = "Your uploaded document includes images. Please try again."

        elif ext == 'docx':
            import docx
            file.seek(0)  # Reset file pointer to the start
            doc = docx.Document(file)
            text = ""
            for para in doc.paragraphs:
                text += para.text
            # Check for images in the document
            if any(run.element.xpath(".//w:drawing") for para in doc.paragraphs for run in para.runs):
                error = "Your uploaded document includes images. Please try again.."

        else:
            error = "Unsupported file type."

    except Exception as e:
        error = f"An error occurred: {str(e)}"

    return text, error


def highlight_text(text1, text2):
    # Use difflib to identify the matching blocks of text
    s = difflib.SequenceMatcher(None, text1.split(), text2.split())
    matching_blocks = s.get_matching_blocks()

    highlighted_text1 = []
    highlighted_text2 = []

    words1 = text1.split()
    words2 = text2.split()

    last_index1 = 0
    last_index2 = 0

    # Highlight matching phrases
    for match in matching_blocks:
        start1, start2, size = match

        # Add the non-matching parts
        highlighted_text1.append(' '.join([f'<span class="different-word">{word}</span>' for word in words1[last_index1:start1]]))
        highlighted_text2.append(' '.join([f'<span class="different-word">{word}</span>' for word in words2[last_index2:start2]]))

        # Add the matching parts
        highlighted_text1.append(' '.join([f'<span class="common-word">{word}</span>' for word in words1[start1:start1 + size]]))
        highlighted_text2.append(' '.join([f'<span class="common-word">{word}</span>' for word in words2[start2:start2 + size]]))

        last_index1 = start1 + size
        last_index2 = start2 + size

    # Add any remaining non-matching parts
    highlighted_text1.append(' '.join([f'<span class="different-word">{word}</span>' for word in words1[last_index1:]]))
    highlighted_text2.append(' '.join([f'<span class="different-word">{word}</span>' for word in words2[last_index2:]]))

    return ' '.join(highlighted_text1), ' '.join(highlighted_text2)

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session or not session['logged_in']:
            flash('Please log in to access this page', 'danger')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Single document comparison routes
# Fix random seed for consistency in langdetect
DetectorFactory.seed = 0

@app.route('/singleComparison', methods=['GET', 'POST'])
@login_required
def singleComparison():
    if request.method == 'POST':
        attempts = 1
        document_name = None
        document_format = None
        success = False
        errors = None
        similarity_results = None
        description = "Single document comparison"

        if 'document' not in request.files:
            errors = "No file part"
            return jsonify({"error": errors}), 400

        file = request.files['document']
        document_name = file.filename
        document_format = file.filename.split('.')[-1]

        if file.filename == '':
            errors = "No selected file"
            return jsonify({"error": errors}), 400

        if file and allowed_file(file.filename):
            try:
                new_document_text, error = read_text_from_file(file, file.filename)
                if error:
                    raise Exception(error)

                # Detect language and ensure it's Somali
                detected_language = detect(new_document_text)
                if detected_language != 'so':  # 'so' is the language code for Somali
                    errors = f"Document is not in Somali language. Detected language: {detected_language}."
                    # Insert error data into the comparison_info table
                    cur = mysql.connection.cursor()
                    cur.execute("""
                        INSERT INTO comparison_info (user_id, comparison_type, document_name, document_format, attempts, success, errors, similarity_results, description)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """, (session['user_id'], 'single', document_name, document_format, attempts, False, errors, similarity_results, description))
                    mysql.connection.commit()
                    cur.close()
                    return jsonify({"error": errors}), 400

                # Process the document comparison
                dataset_paths = [
                    'existing documents/Naftaydaay-Gacalo.txt',
                    'existing documents/Aabe Qani ah Aabe Faqiir ah.txt',
                    'existing documents/NUUN.txt'
                ]
                existing_documents = load_existing_documents(dataset_paths)

                vectorizer = CountVectorizer(binary=True)
                new_doc_vector = vectorizer.fit_transform([new_document_text])

                similarities = []
                for name, text in existing_documents:
                    existing_doc_vector = vectorizer.transform([text])
                    
                    # Calculate Jaccard similarity
                    similarity_score = jaccard_score(
                        new_doc_vector.toarray()[0],
                        existing_doc_vector.toarray()[0],
                        average='binary'
                    )
                    similarities.append((name, similarity_score))

                similarities_percent = [(name, round(similarity * 100, 2)) for name, similarity in similarities]
                similarity_results = max([sim[1] for sim in similarities_percent])
                similarity_threshold = 50.0
                similar_to_any = any(similarity >= similarity_threshold for _, similarity in similarities_percent)

                results = {
                    'similarities': similarities_percent,
                    'similar_to_existing': similar_to_any
                }

                success = True

                # Insert data into the comparison_info table
                cur = mysql.connection.cursor()
                cur.execute("""
                    INSERT INTO comparison_info (user_id, comparison_type, document_name, document_format, attempts, success, errors, similarity_results, description)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (session['user_id'], 'single', document_name, document_format, attempts, success, errors, similarity_results, description))
                mysql.connection.commit()
                cur.close()

                return jsonify(results)
            except Exception as e:
                errors = str(e)
                # Insert error data into the comparison_info table
                cur = mysql.connection.cursor()
                cur.execute("""
                    INSERT INTO comparison_info (user_id, comparison_type, document_name, document_format, attempts, success, errors, similarity_results, description)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (session['user_id'], 'single', document_name, document_format, attempts, False, errors, similarity_results, description))
                mysql.connection.commit()
                cur.close()
                return jsonify({"error": errors}), 400
        else:
            errors = "Invalid file type. Only pdf, docx, and txt are allowed."
            # Insert error data into the comparison_info table
            cur = mysql.connection.cursor()
            cur.execute("""
                INSERT INTO comparison_info (user_id, comparison_type, document_name, document_format, attempts, success, errors, similarity_results, description)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (session['user_id'], 'single', document_name, document_format, attempts, False, errors, similarity_results, description))
            mysql.connection.commit()
            cur.close()
            return jsonify({"error": errors}), 400

    return render_template('single.html')


# Ensure the results are consistent by setting the seed
DetectorFactory.seed = 0

def detect_language(text):
    try:
        return detect(text)
    except:
        return None
@app.route('/doubleComparison', methods=['GET', 'POST'])
@login_required
def doubleComparison():
    if request.method == 'POST':
        file1 = request.files.get('file1')
        file2 = request.files.get('file2')
        attempts = 1
        success = False
        errors = None
        similarity_results = None
        description = "Double document comparison"

        if not file1 or not file2:
            errors = "Both files are required."
            return jsonify({"error": errors}), 400

        if not allowed_file(file1.filename) or not allowed_file(file2.filename):
            errors = "Invalid file type. Only pdf, docx, and txt are allowed."
            return jsonify({"error": errors}), 400

        try:
            text1, error1 = read_text_from_file(file1, file1.filename)
            text2, error2 = read_text_from_file(file2, file2.filename)
            if error1 or error2:
                raise Exception(error1 or error2)

            # Detect language of the documents and ensure they are both Somali
            lang1 = detect_language(text1)
            lang2 = detect_language(text2)
            if lang1 != 'so' or lang2 != 'so':
                raise Exception("Both documents must be in Somali.")

            text1_binary = get_binary_representation(text1, tokenizer_double)
            text2_binary = get_binary_representation(text2, tokenizer_double)

            # Convert sets to binary vectors
            vectorizer = CountVectorizer(binary=True)
            corpus = [' '.join(text1_binary), ' '.join(text2_binary)]
            X = vectorizer.fit_transform(corpus).toarray()

            # Calculate Jaccard similarity
            jaccard_sim = jaccard_score(X[0], X[1])

            similarity_results = round(jaccard_sim * 100, 2)

            threshold = 0.8
            is_similar = jaccard_sim > threshold

            highlighted_text1, highlighted_text2 = highlight_text(text1, text2)

            results = {
                'similarity': similarity_results,
                'is_similar': bool(is_similar),
                'original_text1': text1,
                'original_text2': text2,
                'highlighted_text1': highlighted_text1,
                'highlighted_text2': highlighted_text2
            }

            success = True

            # Insert data into the comparison_info table
            cur = mysql.connection.cursor()
            cur.execute("""
                INSERT INTO comparison_info (user_id, comparison_type, document_name, document_format, attempts, success, errors, similarity_results, description)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (session['user_id'], 'double', f"{file1.filename}, {file2.filename}",
                  f"{file1.filename.split('.')[-1]}, {file2.filename.split('.')[-1]}",
                  attempts, success, errors, similarity_results, description))
            mysql.connection.commit()
            cur.close()

            return jsonify(results)
        except Exception as e:
            errors = str(e)
            # Insert error data into the comparison_info table
            cur = mysql.connection.cursor()
            cur.execute("""
                INSERT INTO comparison_info (user_id, comparison_type, document_name, document_format, attempts, success, errors, similarity_results, description)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (session['user_id'], 'double', f"{file1.filename}, {file2.filename}",
                  f"{file1.filename.split('.')[-1]}, {file2.filename.split('.')[-1]}",
                  attempts, False, errors, similarity_results, description))
            mysql.connection.commit()
            cur.close()
            return jsonify({"error": errors}), 400

    return render_template('double.html')

# Authentication and user management routes
# Define the form
class UserRegisterForm(FlaskForm):
    username = StringField('Username', validators=[
        InputRequired(message='Username is required.'),
        Length(min=5, max=25, message='Username must be between 5 and 25 characters long.'),
        Regexp(r'^[a-zA-Z\s]+$', message='Username must contain only letters and spaces.')
    ])
    email = StringField('Email', validators=[
        InputRequired(message='Email is required.'),
        Email(message='Invalid email address.')
    ])
    password = PasswordField('Password', validators=[
        InputRequired(message='Password is required.'),
        Length(min=8, message='Password must be at least 8 characters long.'),
        EqualTo('confirm', message='Passwords must match.')
    ])
    confirm = PasswordField('Confirm Password', validators=[
        InputRequired(message='Please confirm your password.')
    ])

class UserLoginForm(FlaskForm):
    username = StringField('Username', validators=[
        InputRequired(message='Username is required.'),
        Length(min=5, max=25, message='Username must be between 5 and 25 characters long.')
    ])
    password = PasswordField('Password', validators=[
        InputRequired(message='Password is required.'),
        Length(min=6, max=35, message='Password must be between 6 and 35 characters long.')
    ])

class ChangeUsernameForm(FlaskForm):
    new_username = StringField('New Username', validators=[
        InputRequired(message='New username is required.'),
        Length(min=4, max=25, message='New username must be between 4 and 25 characters long.')
    ])


class ChangeEmailForm(FlaskForm):
    new_email = StringField('New Email', validators=[
        InputRequired(message='New email is required.'),
        Email(message='Invalid email address.')
    ])


class ChangePasswordForm(FlaskForm):
    current_password = PasswordField('Current Password', validators=[InputRequired()])
    new_password = PasswordField('New Password', validators=[
        InputRequired(message='New password is required.'),
        Length(min=8, message='New password must be at least 8 characters long.'),
        EqualTo('confirm_new', message='Passwords must match.')
    ])
    confirm_new = PasswordField('Confirm New Password', validators=[InputRequired(message='Please confirm your new password.')])

 
@app.route('/')
def index():
    if 'logged_in' in session and session['logged_in']:
        return render_template('index.html')
    else:
        flash('Please log in to view this page', 'danger')
        return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    form = UserRegisterForm(request.form)
    if request.method == 'POST' and form.validate_on_submit():
        username = form.username.data
        email = form.email.data
        password = generate_password_hash(form.password.data)  # Hash the password

        # Additional email validation using email_validator
        try:
            # Validate and normalize the email
            valid = validate_email(email)
            email = valid.email  # Use the normalized email
        except EmailNotValidError as e:
            flash(f'Invalid email address: {e}', 'danger')
            return render_template('registration.html', form=form)

        try:
            cur = mysql.connection.cursor()
            # Check if the username or email is already in use
            cur.execute("SELECT * FROM users WHERE username = %s OR email = %s", (username, email))
            existing_user = cur.fetchone()
            if existing_user:
                flash('Username or email is already in use.', 'danger')
                return render_template('registration.html', form=form)
            
            # Insert the new user into the database
            cur.execute("INSERT INTO users(username, email, password) VALUES(%s, %s, %s)", (username, email, password))
            mysql.connection.commit()
            flash('You are now registered and can log in', 'success')
        except Exception as e:
            flash(f'Registration failed: {e}', 'danger')
            print(f"Registration error: {e}")
        finally:
            cur.close()
        return redirect(url_for('login'))

    return render_template('registration.html', form=form)

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = UserLoginForm(request.form)
    if request.method == 'POST' and form.validate():
        username = form.username.data
        password = form.password.data
        cur = mysql.connection.cursor()
        try:
            result = cur.execute("SELECT * FROM users WHERE username = %s", [username])
            if result > 0:
                user = cur.fetchone()
                if check_password_hash(user['password'], password):
                    session['logged_in'] = True
                    session['username'] = user['username']
                    session['user_id'] = user['id']  # Add user ID to session
                    flash('You are now logged in', 'success')
                    return redirect(url_for('index'))
                else:
                    flash('Invalid password', 'danger')
            else:
                flash('Username not found', 'danger')
        except Exception as e:
            flash('An error occurred during login. Please try again later.', 'danger')
            print(f"Login error: {e}")
        finally:
            cur.close()
    return render_template('login.html', form=form)


@app.route('/logout')
def logout():
    session.clear()
    flash('You have successfully logged out', 'success')
    return redirect(url_for('login'))

@app.route('/profile')
def profile():
    if 'user_id' not in session:
        flash('Please log in to view your profile', 'danger')
        return redirect(url_for('login'))

    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM users WHERE id = %s", (session['user_id'],))
    user = cur.fetchone()
    cur.close()

    return render_template('profile.html', user=user)

from datetime import datetime
from flask import render_template, request, session, flash

@app.route('/statistics', methods=['GET', 'POST'])
@login_required
def statistics():
    if 'user_id' not in session:
        flash('Please log in to view your statistics', 'danger')
        return redirect(url_for('login'))

    cur = mysql.connection.cursor()
    try:
        # Fetch the most recent record for the user
        cur.execute("""
            SELECT * FROM comparison_info 
            WHERE user_id = %s 
            ORDER BY uploaded_time DESC 
            LIMIT 1
        """, [session['user_id']])
        recent_stat = cur.fetchone()

        # Fetch all statistics (optional, depending on your needs)
        cur.execute("""
            SELECT * FROM comparison_info 
            WHERE user_id = %s 
            ORDER BY uploaded_time DESC
        """, [session['user_id']])
        all_stats = cur.fetchall()

    except Exception as e:
        flash(f'Error fetching statistics: {e}', 'danger')
        recent_stat = None
        all_stats = []

    finally:
        cur.close()

    return render_template('statistics.html', recent_stat=recent_stat, stats=all_stats)


@app.route('/fetch_data', methods=['GET'])
def fetch_data():
    if 'user_id' not in session:
        return jsonify([]), 401  # Unauthorized

    filter_type = request.args.get('type')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')

    query = "SELECT * FROM comparison_info WHERE user_id = %s"
    params = [session['user_id']]

    if filter_type == 'custom' and start_date and end_date:
        query += " AND uploaded_time BETWEEN %s AND %s"
        params.extend([start_date, end_date])

    cur = mysql.connection.cursor()
    cur.execute(query, params)
    stats = cur.fetchall()
    cur.close()

    return jsonify(stats)
@app.route('/change_username', methods=['GET', 'POST'])
def change_username():
    if 'logged_in' in session and session['logged_in']:
        form = ChangeUsernameForm(request.form)
        if request.method == 'POST' and form.validate():
            new_username = form.new_username.data
            cur = mysql.connection.cursor()
            try:
                # Update the username in the database
                cur.execute("UPDATE users SET username = %s WHERE username = %s", (new_username, session['username']))
                mysql.connection.commit()
                session['username'] = new_username
                flash('Username updated successfully', 'success')
                return redirect(url_for('index'))
            except Exception as e:
                # Handle any database errors
                flash('An error occurred while updating the username. Please try again later.', 'danger')
                print(f"Error updating username: {e}")
            finally:
                cur.close()
        return render_template('change_username.html', form=form)
    else:
        flash('Please log in to view this page', 'danger')
        return redirect(url_for('login'))

@app.route('/change_email', methods=['POST'])
def change_email():
    if 'logged_in' in session and session['logged_in']:
        new_email = request.form.get('new_email')
        
        # Validate email format
        if new_email and '@' in new_email:
            cur = mysql.connection.cursor()
            try:
                # Check if the new email is already in use
                cur.execute("SELECT * FROM users WHERE email = %s", (new_email,))
                existing_user = cur.fetchone()
                if existing_user:
                    flash('Email is already in use.', 'danger')
                else:
                    # Update the email in the database
                    cur.execute("UPDATE users SET email = %s WHERE username = %s", (new_email, session['username']))
                    mysql.connection.commit()
                    flash('Email updated successfully', 'success')
            except Exception as e:
                # Handle any database errors
                flash('An error occurred while updating the email. Please try again later.', 'danger')
                print(f"Error updating email: {e}")
            finally:
                cur.close()
        else:
            flash('Invalid email format.', 'danger')
    else:
        flash('Please log in to view this page.', 'danger')
    
    return redirect(url_for('profile'))

@app.route('/change_password', methods=['POST'])
def change_password():
    if 'logged_in' in session and session['logged_in']:
        current_password = request.form.get('current_password')
        new_password = request.form.get('new_password')
        confirm_new = request.form.get('confirm_new')

        if new_password != confirm_new:
            flash('New passwords do not match', 'danger')
            return redirect(url_for('profile'))

        cur = mysql.connection.cursor()
        result = cur.execute("SELECT * FROM users WHERE username = %s", [session['username']])
        if result > 0:
            user = cur.fetchone()
            if check_password_hash(user['password'], current_password):
                cur.execute("UPDATE users SET password = %s WHERE username = %s", 
                            (generate_password_hash(new_password), session['username']))
                mysql.connection.commit()
                cur.close()
                flash('Password updated successfully', 'success')
            else:
                flash('Current password is incorrect', 'danger')
        else:
            flash('User not found.', 'danger')
        cur.close()
    else:
        flash('Please log in to view this page', 'danger')
    return redirect(url_for('profile'))
     
@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

if __name__ == '__main__':
    app.run(debug=True)
